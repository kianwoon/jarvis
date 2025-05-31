"""
Word document handler for extracting content from DOC/DOCX files
"""
import os
import re
from typing import List, Optional, Dict, Any, Tuple
from datetime import datetime
import hashlib

from app.document_handlers.base import DocumentHandler, ExtractedChunk, ExtractionPreview
from utils.deduplication import hash_text


class WordHandler(DocumentHandler):
    """Handler for Microsoft Word documents (DOC/DOCX)"""
    
    SUPPORTED_EXTENSIONS = ['.doc', '.docx']
    
    def __init__(self):
        super().__init__()
        try:
            import docx
            self.docx = docx
        except ImportError:
            raise ImportError(
                "python-docx is required for Word document processing. "
                "Install it with: pip install python-docx"
            )
    
    def validate(self, file_path: str) -> Tuple[bool, Optional[str]]:
        """Validate if Word document can be processed"""
        if not os.path.exists(file_path):
            return False, "File not found"
            
        file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
        if file_size_mb > self.MAX_FILE_SIZE_MB:
            return False, f"File too large: {file_size_mb:.1f}MB (max: {self.MAX_FILE_SIZE_MB}MB)"
            
        file_ext = os.path.splitext(file_path)[1].lower()
        if file_ext not in self.SUPPORTED_EXTENSIONS:
            return False, f"Unsupported file extension: {file_ext}"
            
        return True, None
    
    def extract(self, file_path: str, options: Dict[str, Any] = None) -> List[ExtractedChunk]:
        """Extract text chunks from Word document"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = os.path.splitext(file_path)[1].lower()
        if file_ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file extension: {file_ext}")
        
        # For .doc files, we might need to convert to .docx first
        if file_ext == '.doc':
            # Note: python-docx doesn't support .doc files directly
            # You might need to use python-docx2txt or convert using LibreOffice
            raise NotImplementedError(
                "Direct .doc file support not implemented. "
                "Please convert to .docx format or use LibreOffice conversion."
            )
        
        # Extract text from DOCX
        try:
            doc = self.docx.Document(file_path)
        except Exception as e:
            raise ValueError(f"Failed to load Word document: {str(e)}")
        
        file_name = os.path.basename(file_path)
        with open(file_path, "rb") as f:
            file_content = f.read()
            file_id = hashlib.sha256(file_content).hexdigest()[:12]
        
        chunks = []
        cumulative_word_count = 0  # Track total words processed for page estimation
        
        # Extract text by sections (headings and paragraphs)
        current_section = []
        current_heading = ""
        chunk_index = 0
        
        for element in doc.paragraphs:
            # Check if this is a heading
            if element.style and element.style.name.startswith('Heading'):
                # Process previous section if it exists
                if current_section:
                    section_chunks, words_processed = self._process_section(
                        current_section, current_heading, chunk_index, 
                        file_id, file_name, cumulative_word_count
                    )
                    chunks.extend(section_chunks)
                    cumulative_word_count += words_processed
                    chunk_index = len(chunks)
                    current_section = []
                
                current_heading = element.text.strip()
            else:
                # Regular paragraph
                para_text = element.text.strip()
                if para_text:
                    current_section.append(para_text)
        
        # Process final section
        if current_section:
            section_chunks, words_processed = self._process_section(
                current_section, current_heading, chunk_index, 
                file_id, file_name, cumulative_word_count
            )
            chunks.extend(section_chunks)
            cumulative_word_count += words_processed
        
        # Extract text from tables
        table_chunks = self._extract_tables(doc, file_id, file_name, len(chunks), cumulative_word_count)
        chunks.extend(table_chunks)
        
        # Extract text from headers/footers if significant
        header_footer_chunks = self._extract_headers_footers(doc, file_id, file_name, len(chunks))
        chunks.extend(header_footer_chunks)
        
        # Filter out low-quality chunks
        quality_chunks = []
        for chunk in chunks:
            quality_score = self.calculate_quality_score(chunk.content)
            if quality_score >= self.MIN_QUALITY_SCORE:
                chunk.metadata['quality_score'] = quality_score
                quality_chunks.append(chunk)
        
        return quality_chunks
    
    def _process_section(self, paragraphs: List[str], heading: str, 
                        start_index: int, file_id: str, file_name: str,
                        cumulative_word_count: int) -> Tuple[List[ExtractedChunk], int]:
        """Process a section of paragraphs into chunks
        Returns: (chunks, total_words_in_section)
        """
        chunks = []
        section_total_words = 0
        
        # Combine paragraphs into reasonable chunks
        current_chunk_text = []
        current_word_count = 0
        
        # Target chunk size (in words)
        target_chunk_size = 300  # Approximately 1500 characters
        
        # Page estimation: ~250-300 words per page is typical for documents
        words_per_page = 250
        
        for para in paragraphs:
            word_count = len(para.split())
            
            # Check if adding this paragraph would exceed target size
            if current_word_count + word_count > target_chunk_size and current_chunk_text:
                # Create chunk
                content = '\n\n'.join(current_chunk_text)
                if heading:
                    content = f"# {heading}\n\n{content}"
                
                # Estimate page number based on cumulative word count
                words_so_far = cumulative_word_count + section_total_words + current_word_count
                estimated_page = max(1, (words_so_far // words_per_page) + 1)
                
                metadata = {
                    'source': file_name.lower(),
                    'doc_type': 'word',
                    'section': heading or f"Section {start_index + len(chunks) + 1}",
                    'chunk_index': start_index + len(chunks),
                    'uploaded_at': datetime.now().isoformat(),
                    'file_id': file_id,
                    'chunk_type': 'text',
                    'word_count': current_word_count,
                    'page': estimated_page
                }
                
                metadata['doc_id'] = f"{file_id}_chunk_{start_index + len(chunks)}"
                metadata['hash'] = hash_text(content)
                
                chunk = ExtractedChunk(
                    content=content,
                    metadata=metadata
                )
                chunks.append(chunk)
                section_total_words += current_word_count
                
                # Reset for next chunk
                current_chunk_text = []
                current_word_count = 0
            
            current_chunk_text.append(para)
            current_word_count += word_count
        
        # Create final chunk if there's remaining content
        if current_chunk_text:
            content = '\n\n'.join(current_chunk_text)
            if heading:
                content = f"# {heading}\n\n{content}"
            
            # Estimate page number based on cumulative word count
            words_so_far = cumulative_word_count + section_total_words + current_word_count
            estimated_page = max(1, (words_so_far // words_per_page) + 1)
            
            metadata = {
                'source': file_name.lower(),
                'doc_type': 'word',
                'section': heading or f"Section {start_index + len(chunks) + 1}",
                'chunk_index': start_index + len(chunks),
                'uploaded_at': datetime.now().isoformat(),
                'file_id': file_id,
                'chunk_type': 'text',
                'word_count': current_word_count,
                'page': estimated_page
            }
            
            metadata['doc_id'] = f"{file_id}_chunk_{start_index + len(chunks)}"
            metadata['hash'] = hash_text(content)
            
            chunk = ExtractedChunk(
                content=content,
                metadata=metadata
            )
            chunks.append(chunk)
            section_total_words += current_word_count
        
        return chunks, section_total_words
    
    def _extract_tables(self, doc, file_id: str, file_name: str, 
                       start_index: int, cumulative_word_count: int) -> List[ExtractedChunk]:
        """Extract content from tables"""
        chunks = []
        
        for table_idx, table in enumerate(doc.tables):
            # Extract table as text
            table_data = []
            headers = []
            
            for row_idx, row in enumerate(table.rows):
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_text.append(cell_text)
                
                if row_idx == 0:
                    headers = row_text
                else:
                    table_data.append(row_text)
            
            if not table_data:
                continue
            
            # Convert table to text representation
            content = self._table_to_text(headers, table_data)
            
            if not content.strip():
                continue
            
            # Estimate page for tables - assume tables take roughly 100 words worth of space
            table_word_equivalent = 100
            words_per_page = 250
            # Add table word equivalent to cumulative count
            words_so_far = cumulative_word_count + (len(chunks) * table_word_equivalent)
            estimated_page = max(1, (words_so_far // words_per_page) + 1)
            
            metadata = {
                'source': file_name.lower(),
                'doc_type': 'word',
                'section': f"Table {table_idx + 1}",
                'chunk_index': start_index + len(chunks),
                'uploaded_at': datetime.now().isoformat(),
                'file_id': file_id,
                'chunk_type': 'table',
                'table_index': table_idx,
                'row_count': len(table_data),
                'column_count': len(headers) if headers else len(table_data[0]) if table_data else 0,
                'page': estimated_page
            }
            
            metadata['doc_id'] = f"{file_id}_table_{table_idx}"
            metadata['hash'] = hash_text(content)
            
            chunk = ExtractedChunk(
                content=content,
                metadata=metadata
            )
            
            # Only include if quality is sufficient
            quality_score = self.calculate_quality_score(content)
            if quality_score >= self.MIN_QUALITY_SCORE:
                chunk.metadata['quality_score'] = quality_score
                chunks.append(chunk)
        
        return chunks
    
    def _table_to_text(self, headers: List[str], data: List[List[str]]) -> str:
        """Convert table data to readable text format"""
        lines = []
        
        # Add headers if present
        if headers and any(h.strip() for h in headers):
            lines.append("Table:")
            lines.append(" | ".join(headers))
            lines.append("-" * 50)
        
        # Add data rows
        for row in data:
            if any(cell.strip() for cell in row):
                lines.append(" | ".join(row))
        
        return "\n".join(lines)
    
    def _extract_headers_footers(self, doc, file_id: str, file_name: str, 
                                start_index: int) -> List[ExtractedChunk]:
        """Extract content from headers and footers if significant"""
        chunks = []
        
        # Note: python-docx has limited support for headers/footers
        # This is a placeholder for future enhancement
        
        return chunks
    
    def get_preview(self, file_path: str, max_chunks: int = 5) -> ExtractionPreview:
        """Get a preview of extraction without full processing"""
        try:
            # Extract a few chunks for preview
            chunks = self.extract(file_path)[:max_chunks]
            
            # Get file metadata
            file_size = os.path.getsize(file_path)
            preview_data = self.extract_preview(file_path)
            
            metadata = {
                'filename': os.path.basename(file_path).lower(),
                'file_size_mb': round(file_size / (1024 * 1024), 2),
                'doc_type': 'word',
                **preview_data.get('metadata', {})
            }
            
            return ExtractionPreview(
                total_chunks=len(chunks),
                sample_chunks=chunks,
                file_metadata=metadata
            )
            
        except Exception as e:
            return ExtractionPreview(
                total_chunks=0,
                sample_chunks=[],
                file_metadata={'error': str(e)},
                warnings=[f"Preview generation failed: {str(e)}"]
            )
    
    def extract_preview(self, file_path: str, max_length: int = 500) -> Dict[str, Any]:
        """Extract preview information from Word document"""
        try:
            doc = self.docx.Document(file_path)
            
            # Get first few paragraphs for preview
            preview_text = []
            char_count = 0
            
            for para in doc.paragraphs[:10]:  # First 10 paragraphs max
                text = para.text.strip()
                if text:
                    if char_count + len(text) > max_length:
                        remaining = max_length - char_count
                        preview_text.append(text[:remaining] + "...")
                        break
                    else:
                        preview_text.append(text)
                        char_count += len(text)
            
            # Get document statistics
            total_paragraphs = len(doc.paragraphs)
            total_tables = len(doc.tables)
            
            # Estimate word count
            word_count = 0
            for para in doc.paragraphs:
                word_count += len(para.text.split())
            
            # Get core properties if available
            properties = {}
            try:
                core_props = doc.core_properties
                if core_props.author:
                    properties['author'] = core_props.author
                if core_props.title:
                    properties['title'] = core_props.title
                if core_props.subject:
                    properties['subject'] = core_props.subject
                if core_props.created:
                    properties['created'] = core_props.created.isoformat()
                if core_props.modified:
                    properties['modified'] = core_props.modified.isoformat()
            except:
                pass
            
            return {
                'preview': '\n\n'.join(preview_text),
                'metadata': {
                    'total_paragraphs': total_paragraphs,
                    'total_tables': total_tables,
                    'estimated_word_count': word_count,
                    'properties': properties
                }
            }
            
        except Exception as e:
            return {
                'preview': f"Error loading preview: {str(e)}",
                'metadata': {}
            }